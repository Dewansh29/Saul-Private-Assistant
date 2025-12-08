import docx

def inspect_template(path):
    doc = docx.Document(path)
    print("--- Paragraphs ---")
    for p in doc.paragraphs:
        if p.text.strip():
            print(p.text)
    
    print("\n--- Tables ---")
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text for cell in row.cells]
            print(" | ".join(row_text))

if __name__ == "__main__":
    inspect_template("backend/template.docx")
