from docx import Document

def inspect_template_content(path):
    try:
        doc = Document(path)
        print(f"--- Template Analysis: {path} ---")
        print(f"Total paragraphs: {len(doc.paragraphs)}")
        print(f"Total tables: {len(doc.tables)}")
        print(f"Total sections: {len(doc.sections)}")
        
        print("\n--- First 10 Paragraphs ---")
        for i, p in enumerate(doc.paragraphs[:10]):
            print(f"{i}: '{p.text}' (style: {p.style.name})")
            
        print("\n--- Section Info ---")
        for i, section in enumerate(doc.sections):
            print(f"Section {i}:")
            print(f"  Header paragraphs: {len(section.header.paragraphs)}")
            print(f"  Footer paragraphs: {len(section.footer.paragraphs)}")
            
    except Exception as e:
        print(f"Error: {e}")

if __name__ == "__main__":
    inspect_template_content("backend/template.docx")
