from docx import Document

def inspect_headers(path):
    try:
        doc = Document(path)
        print("--- Headers ---")
        for section in doc.sections:
            for p in section.header.paragraphs:
                print(f"Header: {p.text}")
        
        print("\n--- Footers ---")
        for section in doc.sections:
            for p in section.footer.paragraphs:
                print(f"Footer: {p.text}")

    except Exception as e:
        print(f"Error: {e}")

if __name__ == "__main__":
    inspect_headers("backend/template.docx")
