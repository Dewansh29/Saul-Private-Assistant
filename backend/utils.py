import fitz
import camelot
import pandas as pd
from typing import List, Dict, Any
import io
import json
import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement

# --- FIX: The Excel processing logic is now much more robust ---
def process_excel_data(dataframes: Dict[str, pd.DataFrame]) -> str:
    """
    Scans through Excel sheets, finds rows with financial keywords across ALL columns,
    and returns a clean string of relevant data for the LLM.
    """
    relevant_data_text = ""
    keywords = ['revenue', 'income', 'expenses', 'profit', 'ebitda', 'assets', 'liabilities', 'equity', 'cash']

    for sheet_name, df in dataframes.items():
        df.columns = [str(col) if 'unnamed' not in str(col).lower() else '' for col in df.columns]
        df_str = df.astype(str).apply(lambda x: x.str.lower())

        relevant_rows = []
        for index, row in df_str.iterrows():
            # CHANGE: Search the ENTIRE row for keywords, not just the first two columns.
            if len(row) > 0 and any(keyword in str(cell) for keyword in keywords for cell in row):
                original_row = df.iloc[index]
                row_text = " | ".join([f"{col}: {val}" for col, val in original_row.items() if pd.notna(val) and str(val).strip() != ''])
                if row_text:
                    relevant_rows.append(row_text)

        if relevant_rows:
            relevant_data_text += f"--- Relevant Data from Sheet: {sheet_name} ---\n"
            relevant_data_text += "\n".join(relevant_rows)
            relevant_data_text += "\n\n"
            
    return relevant_data_text if relevant_data_text else "No financial keywords found in the Excel file."

def extract_tables_from_pdf(pdf_bytes: bytes, page_numbers: List[int]) -> Dict[int, List[pd.DataFrame]]:
    tables_by_page = {}
    if not page_numbers:
        return tables_by_page
    pdf_file = io.BytesIO(pdf_bytes)
    pages_str = ",".join(map(str, page_numbers))
    print(f"Attempting to extract tables from pages using Camelot: {pages_str}")
    try:
        tables = camelot.read_pdf(pdf_file, pages=pages_str, flavor='stream', line_scale=40)
        for table_report in tables:
            page = table_report.page
            if page not in tables_by_page:
                tables_by_page[page] = []
            tables_by_page[page].append(table_report.df)
            print(f"Found a table on page {page} with Camelot.")
        return tables_by_page
    except Exception as e:
        print(f"An error occurred while processing with Camelot: {e}")
        return {}

def extract_text_from_pdf(pdf_bytes: bytes, page_numbers: List[int]) -> str:
    full_text = ""
    pdf_document = fitz.open(stream=pdf_bytes, filetype="pdf")
    for page_num in page_numbers:
        if 1 <= page_num <= len(pdf_document):
            page = pdf_document.load_page(page_num - 1)
            full_text += page.get_text("text") + "\n---\n"
    return full_text

def extract_data_from_excel(file_bytes: bytes, filename: str) -> Dict[str, pd.DataFrame]:
    print(f"--- Attempting to extract data from spreadsheet: {filename} ---")
    dataframes = {}
    try:
        if filename.endswith('.csv'):
            df = pd.read_csv(io.BytesIO(file_bytes))
            dataframes['sheet_1'] = df
        elif filename.endswith(('.xls', '.xlsx')):
            xls = pd.ExcelFile(io.BytesIO(file_bytes), engine='openpyxl')
            for sheet_name in xls.sheet_names:
                df = pd.read_excel(xls, sheet_name=sheet_name)
                dataframes[sheet_name] = df
        print(f"--- Successfully extracted {len(dataframes)} sheet(s) ---")
        return dataframes
    except Exception as e:
        print(f"An error occurred while processing the spreadsheet: {e}")
        return {}

def add_formatted_text(paragraph, text):
    """
    Parse markdown-style formatting and add properly formatted runs to a paragraph.
    Handles **bold** and other common markdown patterns.
    """
    # Pattern to match **bold text**
    pattern = r'\*\*(.*?)\*\*'
    
    last_end = 0
    for match in re.finditer(pattern, text):
        # Add text before the bold part
        if match.start() > last_end:
            paragraph.add_run(text[last_end:match.start()])
        
        # Add bold text
        bold_run = paragraph.add_run(match.group(1))
        bold_run.bold = True
        
        last_end = match.end()
    
    # Add remaining text after last match
    if last_end < len(text):
        paragraph.add_run(text[last_end:])

def create_word_report(report_data: Dict, company_name: str) -> io.BytesIO:
    try:
        # Use the user's template.docx
        print("Loading template.docx for report generation.")
        document = Document('template.docx')
        
        # FIX: Remove empty paragraphs at the beginning to avoid content starting from middle
        # Keep only header/footer content, clear body paragraphs
        paragraphs_to_remove = []
        for paragraph in document.paragraphs:
            # Remove if it's empty or just whitespace
            if not paragraph.text.strip():
                paragraphs_to_remove.append(paragraph)
        
        for paragraph in paragraphs_to_remove:
            p = paragraph._element
            p.getparent().remove(p)
        
        # Add the analysis content to the document
        document.add_heading(f'Subject: Financial Analysis of {company_name}', level=2)
        
        document.add_heading('1. Executive Summary & Final Verdict', level=1)
        summary_text = report_data.get('final_summary', 'No summary was generated.')
        summary_para = document.add_paragraph()
        add_formatted_text(summary_para, summary_text)
        
        document.add_heading('2. Key Financial Metrics Extracted', level=1)
        kpi_data_str = report_data.get('cleaned_data', '{}')
        if kpi_data_str and kpi_data_str != '{}' and "Error" not in kpi_data_str:
            try:
                kpi_data = json.loads(kpi_data_str)
                for key, value in kpi_data.items():
                    # Use 'List Paragraph' style which exists in the template
                    p = document.add_paragraph(style='List Paragraph')
                    p.add_run(f'{key}: ').bold = True
                    p.add_run(str(value) if value is not None else 'N/A')
            except Exception as e:
                print(f"Error formatting KPIs: {e}")
                document.add_paragraph("Could not format the extracted KPI data.")
        else:
            document.add_paragraph("No structured financial KPIs were extracted from the document.")
                
        document.add_heading('3. The Boardroom: Analyst Debate Transcript', level=1)
        debate_entries = report_data.get('debate', [])
        if not debate_entries:
            document.add_paragraph("No debate was generated.")
        else:
            for entry in debate_entries:
                p = document.add_paragraph(style='List Paragraph')
                
                # Parse the entry to separate persona from text
                if ':' in entry:
                    persona, text = entry.split(':', 1)
                    # Add persona in bold
                    p.add_run(f'{persona}: ').bold = True
                    # Add formatted text (handles **bold** markdown)
                    add_formatted_text(p, text.strip())
                else:
                    add_formatted_text(p, entry)
            
        document.add_paragraph("\n" + "---" * 25)
        footer = document.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run("CONFIDENTIAL MEMORANDUM | Generated by the Saul Goodman Financial Strategy Assistant.")
        run.font.size = Pt(10)
        run.italic = True
        
        file_stream = io.BytesIO()
        document.save(file_stream)
        file_stream.seek(0)
        
        return file_stream
        
    except Exception as e:
        print(f"CRITICAL ERROR generating report: {e}")
        import traceback
        traceback.print_exc()
        # Fallback to basic document
        doc = Document()
        doc.add_paragraph(f"Error generating report: {e}")
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        return file_stream