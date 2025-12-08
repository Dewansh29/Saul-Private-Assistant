from docx import Document
import os

try:
    print(f"File size: {os.path.getsize('backend/template.docx')}")
    doc = Document('backend/template.docx')
    print("Successfully loaded template.docx")
    print(f"Paragraph count: {len(doc.paragraphs)}")
    print(f"Table count: {len(doc.tables)}")
except Exception as e:
    print(f"Failed to load: {e}")
